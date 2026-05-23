"""
Бот договоров: загрузка документов сторон, извлечение данных через ИИ,
заполнение шаблона договора, скачивание результата.

POST {action: create_session, contract_type, title, conditions_text}  — создать сессию
POST {action: update_session, session_id, ...}  — обновить условия
POST {action: upload_doc, session_id, doc_type, file_name, file_base64, file_ext}  — загрузить документ
POST {action: fill_contract, session_id}  — заполнить договор через Мелания
POST {action: download_format, session_id, format}  — скачать DOCX/PDF (base64)
GET  /?action=sessions  — список сессий
GET  /?action=session&id=  — одна сессия с документами
GET  /?action=download&id=  — скачать TXT
"""

import json
import os
import base64
import io
import urllib.request
import urllib.error
from datetime import datetime, timezone
import psycopg2
from psycopg2.extras import RealDictCursor
import boto3

SCHEMA = os.environ.get('MAIN_DB_SCHEMA', 't_p71821556_real_estate_catalog_')

CORS = {
    'Access-Control-Allow-Origin': '*',
    'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
    'Access-Control-Allow-Headers': 'Content-Type, X-Auth-Token, X-Authorization, Authorization, X-User-Id, X-Session-Id',
}
ALLOWED_ROLES = ('admin', 'editor', 'manager', 'director', 'broker', 'office_manager')
YANDEX_MODEL = 'yandexgpt/rc'
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_EXTS = {'png', 'jpg', 'jpeg', 'pdf', 'doc', 'docx', 'xls', 'xlsx'}

CONTRACT_TYPES = {
    'lease': 'Договор аренды',
    'sale': 'Договор купли-продажи',
    'agency': 'Агентский договор',
    'service': 'Договор оказания услуг',
    'preliminary': 'Предварительный договор',
    'intent': 'Соглашение о намерениях',
    'custom': 'Произвольный договор',
}


def _get_conn():
    """Подключение к БД с явным search_path."""
    dsn = os.environ['DATABASE_URL']
    return psycopg2.connect(dsn, options=f'-c search_path={SCHEMA},public')


def _ok(body, status=200):
    return {
        'statusCode': status,
        'headers': {**CORS, 'Content-Type': 'application/json'},
        'body': json.dumps(body, ensure_ascii=False, default=str),
    }


def _err(msg, status=400):
    return _ok({'error': msg}, status)


def _q(s, n=500):
    """Экранирование строки для SQL."""
    return (str(s) or '').replace("'", "''")[:n]


def _extract_text_from_file(file_data: bytes, file_ext: str, file_name: str) -> str:
    """Извлекает текст из PDF/DOCX/TXT. Возвращает короткое описание если не удалось."""
    ext = (file_ext or '').lower().lstrip('.')
    try:
        if ext == 'pdf':
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_data))
            pages = []
            # Ограничиваем 30 страницами, чтобы не превысить контекст модели
            for i, page in enumerate(reader.pages[:30]):
                try:
                    pages.append(page.extract_text() or '')
                except Exception:
                    continue
            text = '\n'.join(pages).strip()
            return text[:15000] if text else f'[PDF без распознаваемого текста: {file_name}]'
        elif ext == 'docx':
            from docx import Document
            doc = Document(io.BytesIO(file_data))
            paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            text = '\n'.join(paras).strip()
            return text[:15000] if text else f'[DOCX пустой: {file_name}]'
        elif ext in ('txt',):
            return file_data.decode('utf-8', errors='replace')[:15000]
        else:
            # Изображения/прочее — без OCR
            return f'[Документ "{file_name}" ({ext.upper()}): данные будут учтены при заполнении]'
    except Exception as e:
        return f'[Не удалось извлечь текст из "{file_name}": {type(e).__name__}]'


def _check_rate_limit(cur, user_id: int, action: str = 'fill_contract',
                       max_calls: int = 5, period_minutes: int = 60) -> tuple:
    """Проверяет лимит вызовов. Возвращает (allowed: bool, retry_after_minutes: int)."""
    # INTERVAL нельзя параметризовать через %s — формируем строку безопасно (int)
    minutes = int(period_minutes)
    cur.execute(
        f"SELECT COUNT(*) FROM contract_sessions "
        f"WHERE user_id = %s AND status = 'filled' "
        f"AND updated_at > NOW() - INTERVAL '{minutes} minutes'",
        (user_id,)
    )
    row = cur.fetchone()
    # RealDictCursor вернёт dict, обычный cursor — tuple. Поддерживаем оба.
    if not row:
        count = 0
    else:
        try:
            count = row[0] if not isinstance(row, dict) else list(row.values())[0]
        except Exception:
            count = 0
    if count >= max_calls:
        return (False, minutes)
    return (True, 0)


def _get_user(cur, token):
    if not token:
        return None
    cur.execute(
        f"SELECT u.id, u.role, u.name AS full_name "
        f"FROM {SCHEMA}.sessions s "
        f"JOIN {SCHEMA}.users u ON u.id = s.user_id "
        f"WHERE s.token = %s AND s.expires_at > NOW() AND u.is_active = TRUE",
        (token,)
    )
    return cur.fetchone()


def _s3():
    return boto3.client(
        's3',
        endpoint_url='https://bucket.poehali.dev',
        aws_access_key_id=os.environ['AWS_ACCESS_KEY_ID'],
        aws_secret_access_key=os.environ['AWS_SECRET_ACCESS_KEY'],
    )


def _cdn_url(key):
    return f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{key}"


def _load_yandex_keys(cur) -> tuple:
    """Берём ключи из таблицы settings, fallback на ENV."""
    try:
        cur.execute(f"SELECT yandex_api_key, yandex_folder_id FROM {SCHEMA}.settings ORDER BY id ASC LIMIT 1")
        row = cur.fetchone()
        if row:
            api = (row.get('yandex_api_key') or '').strip() if isinstance(row, dict) else ''
            fld = (row.get('yandex_folder_id') or '').strip() if isinstance(row, dict) else ''
            if api and fld:
                return api, fld
    except Exception:
        pass
    return os.environ.get('YANDEX_API_KEY', '').strip(), os.environ.get('YANDEX_FOLDER_ID', '').strip()


def _yandex_gpt(system: str, user_msg: str, api_key: str = '', folder_id: str = '') -> str:
    """Вызов YandexGPT. Бросает ValueError с понятным сообщением при сбое."""
    api_key = (api_key or os.environ.get('YANDEX_API_KEY', '')).strip()
    folder_id = (folder_id or os.environ.get('YANDEX_FOLDER_ID', '')).strip()
    if not api_key or not folder_id:
        raise ValueError('YandexGPT не настроен: добавьте ключи в Настройки → Интеграции')

    payload = {
        'modelUri': f'gpt://{folder_id}/{YANDEX_MODEL}',
        'completionOptions': {'stream': False, 'temperature': 0.3, 'maxTokens': '6000'},
        'messages': [
            {'role': 'system', 'text': system},
            {'role': 'user', 'text': user_msg},
        ],
    }
    data = json.dumps(payload, ensure_ascii=False).encode('utf-8')
    req = urllib.request.Request(
        'https://llm.api.cloud.yandex.net/foundationModels/v1/completion',
        data=data,
        headers={
            'Content-Type': 'application/json',
            'Authorization': f'Api-Key {api_key}',
            'x-folder-id': folder_id,
        },
        method='POST',
    )
    try:
        with urllib.request.urlopen(req, timeout=55) as r:
            result = json.loads(r.read().decode('utf-8'))
    except urllib.error.HTTPError as e:
        body_text = ''
        try:
            body_text = e.read().decode('utf-8', errors='replace')[:500]
        except Exception:
            pass
        if e.code == 401:
            raise ValueError('YandexGPT отклонил ключ (401). Проверьте API-ключ и Folder ID в Настройки → Интеграции')
        if e.code == 403:
            raise ValueError('YandexGPT: нет доступа (403). Проверьте права сервисного аккаунта на ai.languageModels.user')
        if e.code == 429:
            raise ValueError('YandexGPT: превышен лимит запросов (429). Подождите минуту')
        raise ValueError(f'YandexGPT вернул ошибку {e.code}: {body_text[:200]}')
    except urllib.error.URLError as e:
        raise ValueError(f'Не удалось связаться с YandexGPT: {e.reason}')
    except json.JSONDecodeError:
        raise ValueError('YandexGPT вернул некорректный ответ')

    alts = (result.get('result') or {}).get('alternatives') or []
    if not alts:
        raise ValueError('YandexGPT вернул пустой ответ')
    text = ((alts[0].get('message') or {}).get('text') or '').strip()
    if not text:
        raise ValueError('YandexGPT вернул пустой текст')
    return text


def _fill_via_gpt(session: dict, docs: list, api_key: str = '', folder_id: str = '') -> str:
    now = datetime.now(timezone.utc)
    date_str = now.strftime('%d.%m.%Y')
    contract_type_label = CONTRACT_TYPES.get(session.get('contract_type', ''), 'Договор')
    conditions = session.get('conditions_text') or ''

    doc_parts = []
    for d in docs:
        label_map = {
            'party1': 'Документы Арендодателя (Сторона 1)',
            'party2': 'Документы Арендатора (Сторона 2)',
            'template': 'Шаблон договора',
            'other': 'Прочие документы',
        }
        label = label_map.get(d['doc_type'], d['doc_type'])
        extracted = (d.get('extracted_text') or '').strip()
        if extracted:
            doc_parts.append(f"=== {label}: {d['file_name']} ===\n{extracted}")
        else:
            doc_parts.append(f"=== {label}: {d['file_name']} ===")
    docs_text = '\n\n'.join(doc_parts) if doc_parts else 'Документы не загружены'

    system = f"""Ты — Мелания, опытный юридический помощник. Заполняешь договоры на основе документов сторон.

Сегодня: {date_str}. Тип: {contract_type_label}.

Правила:
1. Используй данные из документов: ФИО, адрес, паспорт, ИНН, ОГРН/ОГРНИП, банк. реквизиты
2. Сторона 1 — Арендодатель, Сторона 2 — Арендатор
3. Включи условия сделки от пользователя
4. Профессиональный юридический язык, дата: {date_str}
5. Если данных нет — [ЗАПОЛНИТЬ: описание]
6. В конце раздел «ПОДПИСИ СТОРОН»

Выведи ТОЛЬКО текст договора без пояснений."""

    user_msg = f"""Тип: {contract_type_label}
Название: {session.get('title', '')}

УСЛОВИЯ СДЕЛКИ:
{conditions if conditions else 'Стандартные условия'}

ДОКУМЕНТЫ:
{docs_text}"""

    return _yandex_gpt(system, user_msg, api_key, folder_id)


def _generate_docx(text: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    for para_text in text.split('\n'):
        stripped = para_text.strip()
        if not stripped:
            continue
        p = doc.add_paragraph()
        if stripped.isupper() or (len(stripped) < 80 and stripped.endswith(':')):
            r = p.add_run(stripped)
            r.bold = True
            r.font.size = Pt(12)
        else:
            r = p.add_run(stripped)
            r.font.size = Pt(12)
        p.paragraph_format.first_line_indent = Cm(1.25)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_pdf(text: str, title: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buf = io.BytesIO()
    body_font = 'Helvetica'
    try:
        resp = urllib.request.urlopen(
            'https://fonts.gstatic.com/s/ptserif/v18/EJRVQgYoZZY2vCFuvDFR.ttf', timeout=5
        )
        font_data = resp.read()
        pdfmetrics.registerFont(TTFont('PTSerif', io.BytesIO(font_data)))
        body_font = 'PTSerif'
    except Exception:
        pass

    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=3*cm, rightMargin=1.5*cm, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('T', parent=styles['Normal'],
        fontName=body_font, fontSize=14, leading=18, alignment=TA_CENTER,
        spaceAfter=12, spaceBefore=6)
    body_style = ParagraphStyle('B', parent=styles['Normal'],
        fontName=body_font, fontSize=11, leading=16, alignment=TA_JUSTIFY,
        spaceAfter=4, firstLineIndent=1.25*cm)
    bold_style = ParagraphStyle('H', parent=styles['Normal'],
        fontName=body_font, fontSize=12, leading=16, spaceAfter=4, spaceBefore=8)

    story = [Paragraph(title.replace('&', '&amp;').replace('<', '&lt;'), title_style),
             Spacer(1, 0.3*cm)]

    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.2*cm))
            continue
        safe = stripped.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        if stripped.isupper() or (len(stripped) < 80 and stripped.endswith(':')):
            story.append(Paragraph(safe, bold_style))
        else:
            story.append(Paragraph(safe, body_style))

    doc.build(story)
    return buf.getvalue()


def handler(event: dict, context) -> dict:
    """Бот договоров: создание, загрузка документов, ИИ-заполнение, скачивание"""
    if event.get('httpMethod') == 'OPTIONS':
        return {'statusCode': 200, 'headers': CORS, 'body': ''}

    qs = event.get('queryStringParameters') or {}
    headers_in = event.get('headers') or {}
    token = headers_in.get('X-Auth-Token') or headers_in.get('x-auth-token') or ''
    method = event.get('httpMethod', 'GET')

    body = {}
    if event.get('body'):
        try:
            body = json.loads(event['body'])
        except Exception:
            pass

    action = qs.get('action') or body.get('action', '')

    conn = _get_conn()
    try:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            user = _get_user(cur, token)
            if not user or user['role'] not in ALLOWED_ROLES:
                return _err('Нет доступа', 403)

            uid = user['id']

            # ── СПИСОК СЕССИЙ ─────────────────────────────────────────
            if action == 'sessions' and method == 'GET':
                cur.execute(
                    "SELECT id, title, contract_type, status, created_at, updated_at "
                    "FROM contract_sessions "
                    "WHERE user_id = %s "
                    "ORDER BY updated_at DESC LIMIT 50",
                    (uid,)
                )
                return _ok({'sessions': [dict(r) for r in cur.fetchall()]})

            # ── ОДНА СЕССИЯ ───────────────────────────────────────────
            if action == 'session' and method == 'GET':
                sid = int(qs.get('id', 0))
                cur.execute(
                    "SELECT * FROM contract_sessions WHERE id = %s AND user_id = %s",
                    (sid, uid)
                )
                s = cur.fetchone()
                if not s:
                    return _err('Сессия не найдена', 404)
                cur.execute(
                    "SELECT id, doc_type, file_name, file_url, file_ext, uploaded_at "
                    "FROM contract_documents WHERE session_id = %s ORDER BY uploaded_at",
                    (sid,)
                )
                docs = [dict(d) for d in cur.fetchall()]
                return _ok({'session': dict(s), 'documents': docs})

            # ── СОЗДАТЬ СЕССИЮ ────────────────────────────────────────
            if action == 'create_session':
                title = _q(body.get('title', 'Новый договор'), 254)
                contract_type = _q(body.get('contract_type', 'custom'), 99)
                conditions = _q(body.get('conditions_text', ''), 9999)
                cur.execute(
                    "INSERT INTO contract_sessions (user_id, title, contract_type, conditions_text) "
                    "VALUES (%s, %s, %s, %s) "
                    "RETURNING id, title, contract_type, status, created_at",
                    (uid, title, contract_type, conditions)
                )
                row = dict(cur.fetchone())
                conn.commit()
                return _ok({'session': row})

            # ── ОБНОВИТЬ УСЛОВИЯ ──────────────────────────────────────
            if action == 'update_session':
                sid = int(body.get('session_id', 0))
                fields = []
                vals = []
                if 'title' in body:
                    fields.append("title = %s"); vals.append(_q(body['title'], 254))
                if 'contract_type' in body:
                    fields.append("contract_type = %s"); vals.append(_q(body['contract_type'], 99))
                if 'conditions_text' in body:
                    fields.append("conditions_text = %s"); vals.append(_q(body['conditions_text'], 9999))
                fields.append("updated_at = NOW()")
                if fields:
                    cur.execute(
                        f"UPDATE contract_sessions SET {', '.join(fields)} "
                        "WHERE id = %s AND user_id = %s",
                        vals + [sid, uid]
                    )
                conn.commit()
                return _ok({'ok': True})

            # ── ЗАГРУЗИТЬ ДОКУМЕНТ ────────────────────────────────────
            if action == 'upload_doc':
                sid = int(body.get('session_id', 0))
                doc_type = _q(body.get('doc_type', 'other'), 49)
                file_name = _q(body.get('file_name', 'doc'), 254)
                file_ext = str(body.get('file_ext', '')).lower().lstrip('.')[:19]
                file_b64 = body.get('file_base64', '')

                if file_ext not in ALLOWED_EXTS:
                    return _err(f'Недопустимый формат .{file_ext}. Разрешены: {", ".join(sorted(ALLOWED_EXTS))}')

                cur.execute(
                    "SELECT id FROM contract_sessions WHERE id = %s AND user_id = %s",
                    (sid, uid)
                )
                if not cur.fetchone():
                    return _err('Сессия не найдена', 404)

                try:
                    file_data = base64.b64decode(file_b64)
                except Exception:
                    return _err('Ошибка декодирования файла')

                if len(file_data) > MAX_FILE_SIZE:
                    return _err('Файл слишком большой (макс. 10 МБ)')

                ct_map = {
                    'pdf': 'application/pdf', 'png': 'image/png',
                    'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
                    'doc': 'application/msword',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'xls': 'application/vnd.ms-excel',
                    'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                }
                content_type = ct_map.get(file_ext, 'application/octet-stream')

                s3 = _s3()
                key = f"contracts/{uid}/{sid}/{file_name}"
                s3.put_object(Bucket='files', Key=key, Body=file_data, ContentType=content_type)
                file_url = _cdn_url(key)

                label_map = {
                    'party1': 'Арендодатель (Сторона 1)',
                    'party2': 'Арендатор (Сторона 2)',
                    'template': 'Шаблон договора',
                    'other': 'Прочие',
                }
                label = label_map.get(doc_type, doc_type)
                # Извлекаем текст из PDF/DOCX для качественного заполнения договора
                doc_text = _extract_text_from_file(file_data, file_ext, file_name)
                extracted = f'[{label}: {file_name}]\n{doc_text}'

                cur.execute(
                    "INSERT INTO contract_documents "
                    "(session_id, doc_type, file_name, file_url, file_ext, extracted_text) "
                    "VALUES (%s, %s, %s, %s, %s, %s) "
                    "RETURNING id, doc_type, file_name, file_url, file_ext",
                    (sid, doc_type, file_name, file_url, file_ext, extracted)
                )
                doc = dict(cur.fetchone())
                cur.execute("UPDATE contract_sessions SET updated_at = NOW() WHERE id = %s", (sid,))
                conn.commit()
                return _ok({'document': doc})

            # ── ЗАПОЛНИТЬ ДОГОВОР ─────────────────────────────────────
            if action == 'fill_contract':
                sid = int(body.get('session_id', 0))

                # Рейт-лимит: не более 5 заполнений в час на пользователя
                allowed, retry_min = _check_rate_limit(cur, uid, 'fill_contract', 5, 60)
                if not allowed:
                    return _err(
                        f'Превышен лимит запросов к ИИ: не более 5 заполнений в час. Попробуйте через {retry_min} минут.',
                        429
                    )

                cur.execute(
                    "SELECT * FROM contract_sessions WHERE id = %s AND user_id = %s",
                    (sid, uid)
                )
                session = cur.fetchone()
                if not session:
                    return _err('Сессия не найдена', 404)

                cur.execute(
                    "SELECT * FROM contract_documents WHERE session_id = %s ORDER BY uploaded_at",
                    (sid,)
                )
                docs = [dict(d) for d in cur.fetchall()]

                api_key, folder_id = _load_yandex_keys(cur)
                if not api_key or not folder_id:
                    return _err('YandexGPT не настроен. Добавьте ключи в Настройки → Интеграции', 503)

                try:
                    filled = _fill_via_gpt(dict(session), docs, api_key, folder_id)
                except ValueError as e:
                    return _err(str(e), 400)
                except Exception as e:
                    return _err(f'Ошибка заполнения договора: {type(e).__name__}: {str(e)[:200]}', 500)

                if not filled or not filled.strip():
                    return _err('ИИ вернул пустой результат. Проверьте загруженные документы и условия сделки.', 400)

                txt_key = f"contracts/{uid}/{sid}/filled_contract.txt"
                s3 = _s3()
                s3.put_object(Bucket='files', Key=txt_key,
                              Body=filled.encode('utf-8'), ContentType='text/plain; charset=utf-8')
                result_url = _cdn_url(txt_key)

                cur.execute(
                    "UPDATE contract_sessions "
                    "SET filled_contract = %s, result_url = %s, status = 'filled', updated_at = NOW() "
                    "WHERE id = %s",
                    (filled[:60000], result_url, sid)
                )
                conn.commit()
                return _ok({'ok': True, 'filled_contract': filled, 'result_url': result_url})

            # ── СКАЧАТЬ TXT ───────────────────────────────────────────
            if action == 'download' and method == 'GET':
                sid = int(qs.get('id', 0))
                cur.execute(
                    "SELECT title, filled_contract FROM contract_sessions "
                    "WHERE id = %s AND user_id = %s",
                    (sid, uid)
                )
                row = cur.fetchone()
                if not row or not row['filled_contract']:
                    return _err('Договор ещё не заполнен', 404)
                return {
                    'statusCode': 200,
                    'headers': {**CORS, 'Content-Type': 'text/plain; charset=utf-8',
                                'Content-Disposition': f'attachment; filename="contract_{sid}.txt"'},
                    'body': row['filled_contract'],
                }

            # ── СКАЧАТЬ DOCX / PDF ────────────────────────────────────
            if action == 'download_format':
                sid = int(body.get('session_id', 0))
                fmt = str(body.get('format', 'docx')).lower()
                if fmt not in ('docx', 'doc', 'pdf'):
                    return _err('Формат должен быть docx, doc или pdf')

                cur.execute(
                    "SELECT title, filled_contract FROM contract_sessions "
                    "WHERE id = %s AND user_id = %s",
                    (sid, uid)
                )
                row = cur.fetchone()
                if not row or not row['filled_contract']:
                    return _err('Договор ещё не заполнен', 404)

                text = row['filled_contract']
                title = row['title'] or f'Договор #{sid}'

                try:
                    if fmt in ('docx', 'doc'):
                        file_bytes = _generate_docx(text, title)
                        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        ext = 'docx'
                    else:
                        file_bytes = _generate_pdf(text, title)
                        content_type = 'application/pdf'
                        ext = 'pdf'
                except Exception as e:
                    return _err(f'Ошибка генерации: {str(e)[:200]}')

                s3 = _s3()
                key = f"contracts/{uid}/{sid}/contract_{sid}.{ext}"
                s3.put_object(Bucket='files', Key=key, Body=file_bytes, ContentType=content_type)
                file_url = _cdn_url(key)

                b64 = base64.b64encode(file_bytes).decode('utf-8')
                safe_title = title.replace(' ', '_').replace('/', '_')[:40]
                return _ok({
                    'ok': True, 'file_base64': b64, 'file_url': file_url,
                    'content_type': content_type, 'filename': f'{safe_title}.{ext}',
                })

            return _err('Неизвестный action', 404)
    finally:
        conn.close()